import streamlit as st
import pandas as pd
from datetime import datetime, date
import os
from PIL import Image
import calendar
import io
from docx import Document
from docx.shared import Inches

# --- Configuración de la página ---
st.set_page_config(page_title="Bitácora Servicios", page_icon="📝", layout="centered")

# --- Configuración de Carpetas ---
CARPETA_DATOS = "datos_bitacora"
CARPETA_IMAGENES = os.path.join(CARPETA_DATOS, "evidencias")
ARCHIVO_CSV = os.path.join(CARPETA_DATOS, "registro_actividades.csv")

os.makedirs(CARPETA_IMAGENES, exist_ok=True)

# --- Funciones de soporte ---
def cargar_datos():
    if os.path.exists(ARCHIVO_CSV):
        try:
            df = pd.read_csv(ARCHIVO_CSV)
            df['Fecha'] = pd.to_datetime(df['Fecha']).dt.date
            return df
        except:
            return pd.DataFrame(columns=["Fecha", "Hora", "Actividad", "RutaImagen"])
    else:
        return pd.DataFrame(columns=["Fecha", "Hora", "Actividad", "RutaImagen"])

def guardar_registro(fecha, hora, actividad, imagen_upload):
    ruta_imagen_final = None
    
    if imagen_upload is not None:
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_imagen = f"evidencia_{timestamp_str}.jpg"
        ruta_imagen_final = os.path.join(CARPETA_IMAGENES, nombre_imagen)
        
        try:
            img = Image.open(imagen_upload)
            img.thumbnail((1000, 1000)) 
            img.save(ruta_imagen_final, quality=85, optimize=True)
        except Exception as e:
            st.error(f"Error al guardar imagen: {e}")

    df = cargar_datos()
    nuevo_registro = pd.DataFrame({
        "Fecha": [fecha],
        "Hora": [hora],
        "Actividad": [actividad],
        "RutaImagen": [ruta_imagen_final if ruta_imagen_final else "Sin evidencia"]
    })
    
    df = pd.concat([df, nuevo_registro], ignore_index=True)
    df.to_csv(ARCHIVO_CSV, index=False)
    return True

def obtener_calendario_dataframe(anio, mes, dias_registrados):
    cal = calendar.monthcalendar(anio, mes)
    cal_visual = []
    for semana in cal:
        semana_visual = []
        for dia in semana:
            if dia == 0:
                semana_visual.append("") 
            else:
                fecha_actual = date(anio, mes, dia)
                if fecha_actual in dias_registrados:
                    semana_visual.append(f"{dia} ✅")
                else:
                    semana_visual.append(f"{dia}")
        cal_visual.append(semana_visual)
    return pd.DataFrame(cal_visual, columns=["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"])

def generar_word(df_filtrado, mes_nombre, anio):
    """Genera un archivo Word en memoria con los datos filtrados"""
    doc = Document()
    
    # Título del documento
    doc.add_heading(f'Informe de Actividades - {mes_nombre} {anio}', 0)
    doc.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y")}')
    
    # Iterar sobre las filas (ordenadas por fecha)
    # Primero las ordenamos por fecha ascendente para el reporte
    df_reporte = df_filtrado.sort_values(by=["Fecha", "Hora"])
    
    for index, row in df_reporte.iterrows():
        fecha_fmt = row['Fecha'].strftime('%d/%m/%Y')
        # Encabezado por cada actividad
        doc.add_heading(f"{fecha_fmt} - {row['Hora']}", level=2)
        
        # Descripción
        p = doc.add_paragraph()
        runner = p.add_run("Descripción: ")
        runner.bold = True
        p.add_run(row['Actividad'])
        
        # Imagen
        ruta_img = row['RutaImagen']
        if ruta_img and ruta_img != "Sin evidencia" and os.path.exists(ruta_img):
            try:
                doc.add_paragraph("Evidencia fotográfica:")
                # Insertar imagen con ancho de 4 pulgadas (para que quepa bien)
                doc.add_picture(ruta_img, width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"[No se pudo cargar la imagen: {e}]")
        
        doc.add_paragraph("-" * 50) # Separador

    # Guardar en memoria (buffer) en lugar de disco
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Interfaz Gráfica ---

st.title("🛠️ Bitácora de Servicios")

tab_registro, tab_calendario, tab_historial = st.tabs(["📝 Nuevo Registro", "📅 Vista Calendario", "📋 Historial y Reporte"])

df = cargar_datos()

# ================= PESTAÑA 1: REGISTRO =================
with tab_registro:
    st.markdown("##### Ingresar Actividad")
    with st.form("formulario_bitacora", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            fecha_input = st.date_input("Fecha", date.today())
        with col2:
            hora_input = st.time_input("Hora", datetime.now().time())
            
        actividad_input = st.text_area("Descripción", height=100)
        imagen_input = st.file_uploader("Evidencia", type=['jpg', 'png', 'jpeg'], label_visibility="collapsed")
        
        submitted = st.form_submit_button("💾 Guardar", type="primary", use_container_width=True)
        
        if submitted:
            if actividad_input.strip() == "":
                st.error("⚠️ Falta descripción.")
            else:
                guardar_registro(fecha_input, hora_input.strftime("%H:%M"), actividad_input, imagen_input)
                st.success("✅ Guardado")
                st.rerun()

# ================= PESTAÑA 2: CALENDARIO =================
with tab_calendario:
    st.subheader("Mapa de cumplimiento")
    col_a, col_b = st.columns(2)
    with col_a:
        anio_ver = st.number_input("Año", value=date.today().year, step=1)
    with col_b:
        mes_ver = st.selectbox("Mes", range(1, 13), index=date.today().month - 1)

    if not df.empty:
        fechas_registradas = set(df['Fecha'].tolist())
    else:
        fechas_registradas = set()

    try:
        st.markdown(f"**{calendar.month_name[mes_ver]} {anio_ver}**")
        df_cal = obtener_calendario_dataframe(anio_ver, mes_ver, fechas_registradas)
        st.dataframe(df_cal, use_container_width=True, height=300)
    except:
        pass

# ================= PESTAÑA 3: HISTORIAL Y REPORTE =================
with tab_historial:
    st.subheader("Generar Informe Mensual")
    
    if df.empty:
        st.info("No hay datos para generar reportes.")
    else:
        # Filtros para el reporte
        col_rep1, col_rep2 = st.columns(2)
        with col_rep1:
            rep_anio = st.number_input("Año Reporte", value=date.today().year, key="rep_anio")
        with col_rep2:
            nombre_meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 
                            7: "Julio", 8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}
            rep_mes_nombre = st.selectbox("Mes Reporte", list(nombre_meses.values()), index=date.today().month - 1)
            # Obtener numero de mes
            rep_mes_num = [k for k, v in nombre_meses.items() if v == rep_mes_nombre][0]

        # Filtrar datos
        df_filtrado = df[
            (pd.to_datetime(df['Fecha']).dt.year == rep_anio) & 
            (pd.to_datetime(df['Fecha']).dt.month == rep_mes_num)
        ]
        
        st.divider()
        
        if df_filtrado.empty:
            st.warning(f"No hay actividades registradas en {rep_mes_nombre} de {rep_anio}.")
        else:
            st.success(f"Se encontraron {len(df_filtrado)} actividades en este periodo.")
            
            # --- BOTÓN DE DESCARGA WORD ---
            archivo_word = generar_word(df_filtrado, rep_mes_nombre, rep_anio)
            
            st.download_button(
                label="📄 Descargar Reporte en Word (.docx)",
                data=archivo_word,
                file_name=f"Reporte_{rep_mes_nombre}_{rep_anio}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
            
            st.markdown("---")
            st.caption("Vista previa de los registros:")
            st.dataframe(df_filtrado[["Fecha", "Hora", "Actividad"]], use_container_width=True)